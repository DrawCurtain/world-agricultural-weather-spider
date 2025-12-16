import os
import datetime
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from downloader import ImageDownloader
from parser import WeatherParser

class DailyWeatherSummary:
    """每日天气数据汇总模块，用于生成今天和前一天的天气对比Word文档"""
    
    def __init__(self):
        self.downloader = ImageDownloader()
        self.parser = WeatherParser()
        self.today = datetime.datetime.now()
        self.yesterday = self.today - datetime.timedelta(days=1)
        self.today_str = self.today.strftime("%Y%m%d")
        self.yesterday_str = self.yesterday.strftime("%Y%m%d")
    
    def download_today_data(self):
        """下载今天的天气数据"""
        print(f"=== 开始下载今天({self.today_str})的天气数据 ===")
        
        # 下载大豆的降水预报
        print("\n下载大豆降水预报数据...")
        pcp_results = self.downloader.download_all_images_by_crop(
            crop_index=1,  # soybeans
            vrbl="pcp",
            nday=15
        )
        
        # 下载大豆的温度预报
        print("\n下载大豆温度预报数据...")
        tmp_results = self.downloader.download_all_images_by_crop(
            crop_index=1,  # soybeans
            vrbl="tmp",
            nday=15
        )
        
        # 统计结果
        pcp_success = sum(1 for success in pcp_results.values() if success)
        tmp_success = sum(1 for success in tmp_results.values() if success)
        
        print(f"\n=== 下载完成 ===")
        print(f"降水数据: 成功 {pcp_success}/{len(pcp_results)}")
        print(f"温度数据: 成功 {tmp_success}/{len(tmp_results)}")
        
        return pcp_results, tmp_results
    
    def find_image_pairs(self, vrbl):
        """查找今天和前一天的图片对
        
        Args:
            vrbl: 天气变量（"pcp"表示降水，"tmp"表示温度）
            
        Returns:
            list: 包含今天和前一天图片路径的元组列表
        """
        image_pairs = []
        
        # 获取大豆的所有地区和子地区
        crop_index = 1  # soybeans
        regions = self.parser.get_regions_by_crop(crop_index)
        
        for region_index, region in enumerate(regions):
            subregions = self.parser.get_subregions_by_crop_and_region(crop_index, region_index)
            
            for subregion_index, subregion in enumerate(subregions):
                # 生成今天和前一天的图片路径（新的目录结构）
                today_path = f"downloads/{vrbl}/{self.today_str}/{vrbl}_soybeans_{regions[region_index]}_{subregion}_forecast.png"
                yesterday_path = f"downloads/{vrbl}/{self.yesterday_str}/{vrbl}_soybeans_{regions[region_index]}_{subregion}_forecast.png"
                
                # 检查图片是否存在
                if os.path.exists(today_path) and os.path.exists(yesterday_path):
                    image_pairs.append((today_path, yesterday_path, regions[region_index], subregion))
                else:
                    print(f"缺少图片对: {today_path} 或 {yesterday_path}")
        
        return image_pairs
    
    def add_image_with_caption(self, doc, image_path, caption):
        """添加图片到Word文档并添加标题"""
        # 添加图片
        paragraph = doc.add_paragraph()
        run = paragraph.add_run()
        run.add_picture(image_path, width=Inches(6))
        
        # 居中图片
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加标题
        caption_paragraph = doc.add_paragraph(caption)
        caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    def create_comparison_document(self, vrbl, image_pairs):
        """创建对比Word文档
        
        Args:
            vrbl: 天气变量（"pcp"表示降水，"tmp"表示温度）
            image_pairs: 图片对列表
            
        Returns:
            str: 生成的Word文档路径
        """
        # 创建文档
        doc = Document()
        
        # 设置文档标题
        title = "大豆作物15天天气预报对比"
        if vrbl == "pcp":
            title += " - 降水"
        else:
            title += " - 温度"
        
        title_paragraph = doc.add_heading(title, level=0)
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加日期信息
        date_paragraph = doc.add_paragraph(f"对比日期: {self.yesterday_str} vs {self.today_str}")
        date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加空行
        doc.add_paragraph()
        
        # 添加每个地区的对比图
        for today_path, yesterday_path, region, subregion in image_pairs:
            # 添加地区标题
            region_title = f"{region} - {subregion}"
            doc.add_heading(region_title, level=2)
            
            # 添加前一天的图片
            self.add_image_with_caption(doc, yesterday_path, f"{self.yesterday_str}")
            
            # 添加今天的图片
            self.add_image_with_caption(doc, today_path, f"{self.today_str}")
            
            # 添加分页符
            doc.add_page_break()
        
        # 保存文档
        doc_path = f"weather_summary_{vrbl}_{self.today_str}.docx"
        doc.save(doc_path)
        
        return doc_path
    
    def run(self):
        """运行完整的汇总流程"""
        # 下载今天的数据
        self.download_today_data()
        
        # 创建降水对比文档
        print(f"\n=== 创建降水对比文档 ===")
        pcp_image_pairs = self.find_image_pairs("pcp")
        if pcp_image_pairs:
            pcp_doc_path = self.create_comparison_document("pcp", pcp_image_pairs)
            print(f"降水对比文档已生成: {pcp_doc_path}")
        else:
            print("没有找到降水图片对，无法生成对比文档")
        
        # 创建温度对比文档
        print(f"\n=== 创建温度对比文档 ===")
        tmp_image_pairs = self.find_image_pairs("tmp")
        if tmp_image_pairs:
            tmp_doc_path = self.create_comparison_document("tmp", tmp_image_pairs)
            print(f"温度对比文档已生成: {tmp_doc_path}")
        else:
            print("没有找到温度图片对，无法生成对比文档")
        
        print(f"\n=== 汇总完成 ===")

if __name__ == "__main__":
    # 检查是否已安装python-docx
    try:
        import docx
    except ImportError:
        print("错误: 未安装python-docx库，请先运行 'pip install python-docx'")
        exit(1)
    
    summary = DailyWeatherSummary()
    summary.run()
